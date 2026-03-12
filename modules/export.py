import os
import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


OUTPUTS_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), "outputs")


def _safe_filename(keyword: str) -> str:
    """Crée un nom de fichier sûr depuis un mot-clé."""
    safe = re.sub(r"[^\w\s-]", "", keyword).strip()
    safe = re.sub(r"[\s]+", "_", safe)
    return safe[:60]


def export_markdown(content: str, keyword: str) -> str:
    """Exporte le contenu en fichier Markdown. Retourne le chemin du fichier."""
    os.makedirs(OUTPUTS_DIR, exist_ok=True)
    filename = f"seo_{_safe_filename(keyword)}.md"
    path = os.path.join(OUTPUTS_DIR, filename)
    with open(path, "w", encoding="utf-8") as f:
        f.write(content)
    return path


def export_word(content: str, keyword: str) -> str:
    """Exporte le contenu en fichier Word (.docx). Retourne le chemin du fichier."""
    os.makedirs(OUTPUTS_DIR, exist_ok=True)
    filename = f"seo_{_safe_filename(keyword)}.docx"
    path = os.path.join(OUTPUTS_DIR, filename)

    doc = Document()

    # Style de base
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    lines = content.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        if line.startswith("# "):
            p = doc.add_heading(line[2:].strip(), level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("#### "):
            doc.add_heading(line[5:].strip(), level=4)
        elif line.startswith("---"):
            doc.add_paragraph("─" * 50)
        elif line.startswith("- ") or line.startswith("* "):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
        elif re.match(r"^\d+\. ", line):
            doc.add_paragraph(re.sub(r"^\d+\. ", "", line), style="List Number")
        elif line.startswith("|"):
            # Tableau Markdown — converti en texte simple pour Word
            cells = [c.strip() for c in line.strip("|").split("|")]
            doc.add_paragraph(" | ".join(cells))
        elif line.startswith("**") and line.endswith("**"):
            p = doc.add_paragraph()
            run = p.add_run(line.strip("*"))
            run.bold = True
        elif line.strip():
            doc.add_paragraph(line)
        else:
            # Ligne vide → espace
            if i > 0 and lines[i - 1].strip():
                doc.add_paragraph("")

        i += 1

    doc.save(path)
    return path


def build_full_export(
    keyword: str,
    content_gap: str,
    pain_points: str,
    statistics: str,
    seo_brief: str,
    introduction: str,
    article: str,
) -> str:
    """Assemble tout le contenu en un seul document Markdown."""
    sections = [
        f"# Article SEO — {keyword}\n",
        "---\n",
        "## Analyse Content Gap\n",
        content_gap,
        "\n---\n",
        "## Points de Douleur Clients\n",
        pain_points,
        "\n---\n",
        "## Statistiques & Data-mining\n",
        statistics,
        "\n---\n",
        "## Brief SEO + Sémantique\n",
        seo_brief,
        "\n---\n",
        "## Article Complet\n",
        f"### Introduction\n\n{introduction}\n",
        article,
    ]
    return "\n".join(sections)
